#!/usr/bin/env python3
"""Build an editable, styled DOCX counterpart of the final MPP PDF."""

from __future__ import annotations

import csv
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "07-manuscript"
SUBMISSION_DIR = ROOT / "09-submission"
SOURCE_MD = MANUSCRIPT_DIR / "manuscript.md"
DOCX_SOURCE = SUBMISSION_DIR / "ChronosAudit_TemporalClone_MPP_docx-source.md"
OUTPUT_DOCX = SUBMISSION_DIR / "ChronosAudit_TemporalClone_MPP_draft.docx"
PANDOC = Path("/opt/homebrew/bin/pandoc")


def run(argv: list[str], cwd: Path | None = None) -> None:
    subprocess.run(argv, cwd=cwd, check=True)


def diagram_svg(path: Path, *, title: str, labels: list[str], colors: list[str], width: int = 1400) -> None:
    margin = 35
    gap = 22
    height = 230
    box_width = (width - 2 * margin - gap * (len(labels) - 1)) / len(labels)
    elements = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="white"/>',
        f'<text x="{width/2}" y="34" text-anchor="middle" font-family="Arial, sans-serif" font-size="25" font-weight="700" fill="#172033">{title}</text>',
        '<defs><marker id="arrow" markerWidth="10" markerHeight="8" refX="9" refY="4" orient="auto"><path d="M0,0 L10,4 L0,8 z" fill="#46526A"/></marker></defs>',
    ]
    y = 72
    box_height = 105
    for index, (label, color) in enumerate(zip(labels, colors)):
        x = margin + index * (box_width + gap)
        if index:
            previous_right = x - gap
            elements.append(f'<line x1="{previous_right + 4}" y1="{y + box_height/2}" x2="{x - 6}" y2="{y + box_height/2}" stroke="#46526A" stroke-width="3" marker-end="url(#arrow)"/>')
        elements.append(f'<rect x="{x:.1f}" y="{y}" width="{box_width:.1f}" height="{box_height}" rx="12" fill="{color}" stroke="#334155" stroke-width="2"/>')
        words = label.split()
        lines: list[str] = []
        current = ""
        for word in words:
            trial = f"{current} {word}".strip()
            if len(trial) > 17 and current:
                lines.append(current)
                current = word
            else:
                current = trial
        if current:
            lines.append(current)
        start_y = y + box_height / 2 - (len(lines) - 1) * 11
        for offset, line in enumerate(lines):
            elements.append(f'<text x="{x + box_width/2:.1f}" y="{start_y + offset*23:.1f}" text-anchor="middle" font-family="Arial, sans-serif" font-size="18" fill="#172033">{line}</text>')
    elements.append('</svg>')
    path.write_text("\n".join(elements) + "\n", encoding="utf-8")


def make_diagrams() -> tuple[Path, Path]:
    architecture_svg = SUBMISSION_DIR / "docx-architecture.svg"
    timeline_svg = SUBMISSION_DIR / "docx-timeline.svg"
    architecture_png = SUBMISSION_DIR / "docx-architecture.png"
    timeline_png = SUBMISSION_DIR / "docx-timeline.png"
    diagram_svg(
        architecture_svg,
        title="Analysis architecture",
        labels=[
            "417-case index",
            "Closed historical envelopes",
            "Fail-closed evidence gate",
            "Canonical analysis rows",
            "Temporal identity proxy audits",
            "T1-T8 figures manuscript",
            "Human and independent gates",
        ],
        colors=["#EAF3FF", "#EAF3FF", "#FFF3D6", "#E8F7EE", "#EEF0FF", "#EEF0FF", "#FDECEC"],
    )
    diagram_svg(
        timeline_svg,
        title="Minimum Publishable Prototype execution timeline",
        labels=[
            "Aug 28 scope freeze",
            "Aug 28-29 analysis tests",
            "Aug 29 artifacts manuscript",
            "Aug 30 PDF QA",
            "Aug 31-Sep 2 independent review",
            "Sep 3 human upload",
        ],
        colors=["#E8F7EE", "#E8F7EE", "#EEF0FF", "#FFF3D6", "#FFF3D6", "#FDECEC"],
    )
    run(["sips", "-s", "format", "png", str(architecture_svg), "--out", str(architecture_png)])
    run(["sips", "-s", "format", "png", str(timeline_svg), "--out", str(timeline_png)])
    return architecture_png, timeline_png


def identity_table_markdown() -> str:
    path = ROOT / "06-visuals/tables/t5-ablation-or-mechanism.csv"
    with path.open(encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    lines = [
        "| Identity abstraction | Unique | Groups | Rows | Cross-address groups | Cross-chain groups | Max size |",
        "|:---|---:|---:|---:|---:|---:|---:|",
    ]
    labels = {
        "chain_address": "Chain-address",
        "address_only": "Address only",
        "runtime_bytecode": "Exact runtime",
        "metadata_stripped_bytecode": "Metadata-stripped runtime",
    }
    for row in rows:
        lines.append(
            "| {label} | {unique} | {groups} | {duplicate_rows} | {cross_address} | {cross_chain} | {maximum} |".format(
                label=labels[row["identity_abstraction"]],
                unique=row["unique_identities"],
                groups=row["duplicate_groups"],
                duplicate_rows=row["duplicate_rows"],
                cross_address=row["cross_address_duplicate_groups"],
                cross_chain=row["cross_chain_duplicate_groups"],
                maximum=row["maximum_group_size"],
            )
        )
    lines.append("\n*Identity-abstraction mechanism audit (T5).*\n")
    return "\n".join(lines)


def build_docx_source(architecture_png: Path, timeline_png: Path) -> None:
    text = SOURCE_MD.read_text(encoding="utf-8")
    header = """---
title: "Auditing Temporal and Identity Dependence in SCONE-bench"
subtitle: "Dual-Provider Historical State Reconstruction of 417 Smart-Contract Incidents"
author: "Author details pending accountable-author confirmation"
date: "Draft generated 29 August 2026"
lang: en-US
---

"""
    text = re.sub(
        r"^# Auditing Temporal and Identity Dependence in SCONE-bench\n\n## Dual-Provider Historical State Reconstruction of 417 Smart-Contract Incidents\n\n\*\*Author:\*\*.*?\n\n",
        "",
        text,
        count=1,
        flags=re.DOTALL,
    )
    architecture_block = re.compile(r'<figure id="fig:architecture".*?</figure>', re.DOTALL)
    text = architecture_block.sub(
        f'![Analysis architecture]({architecture_png.name})\n\n*Figure 1. Analysis architecture. Arrows represent data flow, not causality; editable Mermaid source is retained in the artifact bundle.*',
        text,
    )
    empty_identity = '<div class="table*">\n\n</div>\n\n<figure id="fig:identity"'
    text = text.replace(empty_identity, identity_table_markdown() + '\n<figure id="fig:identity"')
    text = text.replace(
        '<div class="center">\n\n</div>',
        f'![Execution timeline]({timeline_png.name})\n\n*Execution timeline. Dates after local completion are targets for independent and accountable-human gates.*',
    )
    crossrefs = {
        "tab:priorart": "1",
        "tab:cohort": "2",
        "tab:primary": "3",
        "tab:identity": "4",
        "fig:temporal": "2",
        "fig:identity": "4",
    }

    def clean_link(match: re.Match[str]) -> str:
        content = re.sub(r"<[^>]+>", "", match.group(0))
        for key, value in crossrefs.items():
            content = content.replace(f"[{key}]", value)
        return content

    text = re.sub(r'<a href="#[^"]+".*?</a>', clean_link, text, flags=re.DOTALL)
    text = re.sub(r'<img src="([^"]+)"\s*/>', r'![Quantitative figure](\1)', text)
    text = re.sub(r'</?figure[^>]*>', '', text)
    text = re.sub(r'</?figcaption>', '', text)
    text = re.sub(r'</?div[^>]*>', '', text)
    text = re.sub(r'<span class="math inline"><em>N</em>.*?417</span>', '*N* = 417', text)
    text = re.sub(r'\$`([^`]+)`\$', r'$\1$', text)
    text = text.replace(' {#tab:cohort}', '')
    text = re.sub(r'\n{3,}', '\n\n', text)
    DOCX_SOURCE.write_text(header + text.strip() + "\n", encoding="utf-8")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def style_docx(path: Path) -> None:
    document = Document(path)
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(24)
    section.right_margin = Mm(24)
    section.gutter = Mm(0)
    section.header_distance = Mm(9)
    section.footer_distance = Mm(9)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size, color in (
        ("Title", 19, "172033"),
        ("Subtitle", 12, "46526A"),
        ("Heading 1", 15, "172033"),
        ("Heading 2", 12, "1F4E79"),
        ("Heading 3", 11, "1F4E79"),
    ):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Aptos Display" if style_name != "Normal" else "Times New Roman"
            style.font.size = Pt(size)
            style.font.color.rgb = __import__("docx").shared.RGBColor.from_string(color)
    document.styles["Title"].paragraph_format.space_after = Pt(6)
    document.styles["Heading 1"].paragraph_format.space_before = Pt(12)
    document.styles["Heading 1"].paragraph_format.space_after = Pt(5)

    for table in document.tables:
        table.style = "Table"
        table.autofit = True
        if table.rows:
            header_row = table.rows[0]
            tr_pr = header_row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
            for cell in header_row.cells:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:fill"), "D9EAF7")
                cell._tc.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(8.5)
        for row in table.rows[1:]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)

    max_width = Inches(6.35)
    for shape in document.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

    header = section.header.paragraphs[0]
    header.text = "ChronosAudit TemporalClone MPP — submission-ready draft"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.name = "Aptos"
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = __import__("docx").shared.RGBColor.from_string("64748B")
    add_page_field(section.footer.paragraphs[0])

    properties = document.core_properties
    properties.title = "Auditing Temporal and Identity Dependence in SCONE-bench"
    properties.subject = "ChronosAudit TemporalClone minimum publishable prototype"
    properties.keywords = "smart contracts; benchmark validity; historical state; code reuse; proxy contracts; reproducibility"
    properties.comments = "Editable DOCX counterpart of the mechanically checked PDF. Accountable-human submission gates remain open."
    document.save(path)


def main() -> int:
    architecture_png, timeline_png = make_diagrams()
    build_docx_source(architecture_png, timeline_png)
    run(
        [
            str(PANDOC),
            str(DOCX_SOURCE),
            "--from=gfm+raw_html+tex_math_dollars",
            "--to=docx",
            f"--resource-path={MANUSCRIPT_DIR}:{SUBMISSION_DIR}",
            "--standalone",
            "--output",
            str(OUTPUT_DOCX),
        ],
        cwd=MANUSCRIPT_DIR,
    )
    style_docx(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
